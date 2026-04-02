#!/usr/bin/env python3
"""Parse Word (.doc, .docx) - extract paragraphs and tables. Input: path or content_base64 (from fetch).

Skill subprocess: read one JSON object from stdin (parameters may be nested
under ``arguments`` or passed flat). Write one JSON object to stdout.
"""
import base64
import json
import subprocess
import sys
import tempfile
from io import BytesIO
from pathlib import Path

from common.path_utils import ensure_in_workspace as _ensure_in_workspace

try:
    from constants import (
        OBSERVATION_PREVIEW_LEN,
        WORD_MAX_PARAGRAPHS,
        WORD_MAX_TABLES,
    )
except ImportError:
    OBSERVATION_PREVIEW_LEN = 500
    WORD_MAX_PARAGRAPHS = 10000
    WORD_MAX_TABLES = 500

_DOC_MAGIC = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"
_DOCX_MAGIC = b"PK"


def _detect_word_format(data: bytes) -> str:
    """Return 'doc' or 'docx' from magic bytes."""
    if len(data) >= 8 and data[:8] == _DOC_MAGIC:
        return "doc"
    if len(data) >= 2 and data[:2] == _DOCX_MAGIC:
        return "docx"
    return "docx"  # default guess


def _doc_to_docx_bytes(data: bytes) -> bytes:
    """Convert legacy .doc bytes to .docx using LibreOffice. Requires soffice in PATH."""
    with tempfile.TemporaryDirectory(prefix="word_") as tmp:
        doc_path = Path(tmp) / "input.doc"
        doc_path.write_bytes(data)
        try:
            subprocess.run(
                ["soffice", "--headless", "--convert-to", "docx", "--outdir", tmp, str(doc_path)],
                capture_output=True,
                timeout=60,
                check=True,
            )
        except (subprocess.CalledProcessError, FileNotFoundError, subprocess.TimeoutExpired) as e:
            raise RuntimeError(
                "Legacy .doc conversion requires LibreOffice (soffice). Install it and ensure soffice is in PATH."
            ) from e
        docx_path = Path(tmp) / "input.docx"
        if not docx_path.exists():
            raise RuntimeError("LibreOffice did not produce output.docx")
        return docx_path.read_bytes()


def _parse_word_bytes(data: bytes, path_hint: str | None = None) -> tuple[list[str], list[list[list[str]]]]:
    """Parse .doc or .docx bytes. path_hint used to infer format when ambiguous."""
    if path_hint:
        suffix = Path(path_hint).suffix.lower()
        fmt = "doc" if suffix == ".doc" else "docx"
    else:
        fmt = _detect_word_format(data)
    if fmt == "doc":
        data = _doc_to_docx_bytes(data)
    return _parse_docx_bytes(data)


def _parse_docx_bytes(data: bytes) -> tuple[list[str], list[list[list[str]]]]:
    from docx import Document

    doc = Document(BytesIO(data))

    paragraphs: list[str] = []
    for p in doc.paragraphs:
        if len(paragraphs) >= WORD_MAX_PARAGRAPHS:
            break
        text = (p.text or "").strip()
        if text:
            paragraphs.append(text)

    tables: list[list[list[str]]] = []
    for tbl in doc.tables:
        if len(tables) >= WORD_MAX_TABLES:
            break
        rows: list[list[str]] = []
        for row in tbl.rows:
            cells = [str(cell.text or "").strip() for cell in row.cells]
            rows.append(cells)
        tables.append(rows)

    return paragraphs, tables


def _build_full_text(paragraphs: list[str], tables: list[list[list[str]]]) -> str:
    """Build full text from paragraphs and tables."""
    parts = list(paragraphs)
    for tbl in tables:
        for row in tbl:
            parts.append(" | ".join(row))
    return "\n".join(parts)


def _maybe_write_output(
    workspace_root: Path, output_path: str | None, full_text: str
) -> bool:
    """Write full_text to output_path if set. Returns False on error (caller should return)."""
    if not output_path:
        return True
    out_target = (workspace_root / output_path).resolve()
    if not _ensure_in_workspace(workspace_root, out_target):
        print(json.dumps({"error": "output_path cannot escape workspace"}))
        return False
    out_target.parent.mkdir(parents=True, exist_ok=True)
    out_target.write_text(full_text, encoding="utf-8")
    return True


def _observation_preview(full_text: str) -> str:
    """Truncate for observation preview."""
    if len(full_text) > OBSERVATION_PREVIEW_LEN:
        return full_text[:OBSERVATION_PREVIEW_LEN] + "..."
    return full_text or "(no extractable content)"


def _build_result(
    paragraphs: list[str],
    tables: list[list[list[str]]],
    obs: str,
    output_path: str | None,
) -> dict:
    """Build result dict."""
    result: dict = {
        "paragraphs": paragraphs,
        "tables": tables,
        "observation": obs,
    }
    if output_path:
        result["output_path"] = output_path
        result["written"] = True
    return result


def main() -> None:
    """Run the skill entrypoint (stdin JSON → stdout JSON)."""
    params = json.loads(sys.stdin.read())
    args = params.get("arguments", params)
    workspace_root = Path(params.get("workspace_root", ""))

    path = (args.get("path") or "").strip()
    content_base64 = (args.get("content_base64") or "").strip()

    if path and content_base64:
        print(json.dumps({"error": "Provide path OR content_base64, not both"}))
        return
    if not path and not content_base64:
        print(json.dumps({"error": "path or content_base64 is required"}))
        return

    try:
        if content_base64:
            data = base64.standard_b64decode(content_base64)
        else:
            target = (workspace_root / path).resolve()
            if not target.exists() or not target.is_file():
                print(json.dumps({"error": f"File not found: {path}"}))
                return
            if not _ensure_in_workspace(workspace_root, target):
                print(json.dumps({"error": "Path cannot escape workspace"}))
                return
            data = target.read_bytes()

        path_hint = path if path else None
        paragraphs, tables = _parse_word_bytes(data, path_hint)
        full_text = _build_full_text(paragraphs, tables)
        output_path = (args.get("output_path") or "").strip() or None
        if not _maybe_write_output(workspace_root, output_path, full_text):
            return
        obs = _observation_preview(full_text)
        result = _build_result(paragraphs, tables, obs, output_path)
        print(json.dumps(result, ensure_ascii=False))
    except Exception as e:
        print(json.dumps({"error": str(e)}))


if __name__ == "__main__":
    main()
